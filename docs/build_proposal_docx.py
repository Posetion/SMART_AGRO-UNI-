# -*- coding: utf-8 -*-
"""Build the Smart Agro Community competition proposal as a Word document."""
from pathlib import Path

from docx import Document
import re

from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.oxml import OxmlElement
from docx.oxml.ns import qn, nsmap
from docx.shared import Cm, Inches, Pt, RGBColor

GREEN = RGBColor(0x1B, 0x43, 0x32)
GREEN2 = RGBColor(0x2E, 0x7D, 0x32)
INK = RGBColor(0x21, 0x21, 0x21)
MUTED = RGBColor(0x5A, 0x65, 0x5C)
WHITE = RGBColor(0xFF, 0xFF, 0xFF)
ROW_ALT = "F3F8F4"
HEADER_BG = "1B4332"
DOCS = Path(r"d:\SMART-AGRO\docs")
OUT = DOCS / "Smart-Agro-Community-Competition-Proposal.docx"
SRS_MD = DOCS / "SRS-Smart-Agro-Community.md"
SRS_OUT = DOCS / "SRS-Smart-Agro-Community.docx"
EVAL_IMG = DOCS / "ucs-meiktila-evaluation-criteria.png"


def shade(cell, hex_color: str) -> None:
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), hex_color)
    shd.set(qn("w:val"), "clear")
    tcPr.append(shd)


def set_run_font(run, name="Calibri", size=11, bold=False, italic=False, color=INK):
    run.font.name = name
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Myanmar Text")
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic
    run.font.color.rgb = color


def add_bottom_border(paragraph, color="2E7D32", sz="12"):
    pPr = paragraph._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), sz)
    bottom.set(qn("w:space"), "8")
    bottom.set(qn("w:color"), color)
    pBdr.append(bottom)
    pPr.append(pBdr)


def para(doc, text, *, size=11, bold=False, italic=False, color=INK, space_after=8, space_before=0, align=None, name="Calibri"):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(space_after)
    p.paragraph_format.space_before = Pt(space_before)
    p.paragraph_format.line_spacing = 1.15
    if align:
        p.alignment = align
    run = p.add_run(text)
    set_run_font(run, name=name, size=size, bold=bold, italic=italic, color=color)
    return p


def mixed_para(doc, parts, *, space_after=8, space_before=0, align=None, size=11):
    """parts: list of (text, bold, italic, color?)"""
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(space_after)
    p.paragraph_format.space_before = Pt(space_before)
    p.paragraph_format.line_spacing = 1.15
    if align:
        p.alignment = align
    for part in parts:
        text, bold, italic = part[0], part[1], part[2]
        color = part[3] if len(part) > 3 else INK
        sz = part[4] if len(part) > 4 else size
        run = p.add_run(text)
        set_run_font(run, size=sz, bold=bold, italic=italic, color=color)
    return p


def heading1(doc, text):
    p = para(doc, text, size=16, bold=True, color=GREEN, space_before=16, space_after=8)
    add_bottom_border(p)
    return p


def heading2(doc, text):
    return para(doc, text, size=13, bold=True, color=GREEN2, space_before=12, space_after=6)


def bullet(doc, text, *, bold_lead=None):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.line_spacing = 1.15
    if bold_lead:
        r = p.add_run(bold_lead)
        set_run_font(r, size=11, bold=True, color=INK)
        r2 = p.add_run(text)
        set_run_font(r2, size=11, color=INK)
    else:
        r = p.add_run(text)
        set_run_font(r, size=11, color=INK)
    return p


def add_table(doc, headers, rows):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = "Table Grid"
    table.autofit = True
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = ""
        p = cell.paragraphs[0]
        run = p.add_run(h)
        set_run_font(run, size=10, bold=True, color=WHITE)
        shade(cell, HEADER_BG)
    for r_i, row in enumerate(rows):
        for c_i, val in enumerate(row):
            cell = table.rows[r_i + 1].cells[c_i]
            cell.text = ""
            p = cell.paragraphs[0]
            run = p.add_run(val)
            set_run_font(run, size=10, color=INK)
            if r_i % 2 == 1:
                shade(cell, ROW_ALT)
    spacer = doc.add_paragraph()
    spacer.paragraph_format.space_after = Pt(6)
    return table


def set_cell_widths(table, widths_cm):
    for row in table.rows:
        for i, w in enumerate(widths_cm):
            row.cells[i].width = Cm(w)


def page_break(doc):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.space_before = Pt(0)
    run = p.add_run()
    run.add_break(WD_BREAK.PAGE)


def save_docx(doc, path: Path) -> Path:
    path.parent.mkdir(parents=True, exist_ok=True)
    candidates = [
        path,
        path.with_name(path.stem + "-updated.docx"),
        path.with_name(path.stem + "-with-srs.docx"),
        path.with_name(path.stem + "-v2.docx"),
    ]
    last_err = None
    for dest in candidates:
        try:
            doc.save(str(dest))
            if dest != path:
                print("LOCKED:", path)
            return dest
        except PermissionError as err:
            last_err = err
            continue
    raise last_err


def add_picture(doc, image_path: Path, width_in=6.3):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(6)
    run = p.add_run()
    run.add_picture(str(image_path), width=Inches(width_in))
    return p


def strip_md_links(text: str) -> str:
    return re.sub(r"\[([^\]]+)\]\([^)]+\)", r"\1", text)


def add_md_runs(paragraph, text: str, *, size=11, color=INK):
    text = strip_md_links(text)
    parts = re.split(r"(\*\*[^*]+\*\*|`[^`]+`)", text)
    for part in parts:
        if not part:
            continue
        if part.startswith("**") and part.endswith("**") and len(part) >= 4:
            run = paragraph.add_run(part[2:-2])
            set_run_font(run, size=size, bold=True, color=color)
        elif part.startswith("`") and part.endswith("`") and len(part) >= 2:
            run = paragraph.add_run(part[1:-1])
            set_run_font(run, name="Consolas", size=size, color=color)
        else:
            run = paragraph.add_run(part.replace("*", ""))
            set_run_font(run, size=size, color=color)


def add_md_table(doc, rows):
    if not rows:
        return None
    cols = max(len(r) for r in rows)
    norm = [r + [""] * (cols - len(r)) for r in rows]
    headers, body = norm[0], norm[1:]
    return add_table(doc, headers, body)


def append_markdown(doc, md_text: str) -> None:
    """Render GitHub-flavoured markdown (headings, tables, lists, code) into the doc."""
    lines = md_text.replace("\r\n", "\n").split("\n")
    i = 0
    in_code = False
    code_buf = []

    def flush_code():
        nonlocal code_buf
        if not code_buf:
            return
        block = "\n".join(code_buf).rstrip()
        for cline in block.split("\n") or [""]:
            p = doc.add_paragraph()
            p.paragraph_format.space_after = Pt(0)
            p.paragraph_format.space_before = Pt(0)
            p.paragraph_format.line_spacing = 1.0
            run = p.add_run(cline if cline else " ")
            set_run_font(run, name="Consolas", size=9, color=GREEN)
        spacer = doc.add_paragraph()
        spacer.paragraph_format.space_after = Pt(8)
        code_buf = []

    def parse_row(line: str):
        return [c.replace("**", "").replace("`", "").strip() for c in line.strip().strip("|").split("|")]

    def is_sep(line: str) -> bool:
        s = line.strip().strip("|").replace(" ", "")
        return bool(s) and all(ch in "-:" for ch in s.replace("|", ""))

    while i < len(lines):
        line = lines[i]
        if line.startswith("```"):
            if in_code:
                flush_code()
                in_code = False
            else:
                in_code = True
                code_buf = []
            i += 1
            continue
        if in_code:
            code_buf.append(line)
            i += 1
            continue
        if line.strip() in ("---", "***"):
            i += 1
            continue
        if line.strip() == "":
            i += 1
            continue
        if line.startswith("|") and i + 1 < len(lines) and is_sep(lines[i + 1]):
            rows = [parse_row(line)]
            i += 2
            while i < len(lines) and lines[i].strip().startswith("|"):
                if not is_sep(lines[i]):
                    rows.append(parse_row(lines[i]))
                i += 1
            add_md_table(doc, rows)
            continue
        if line.startswith("#"):
            hashes = len(line) - len(line.lstrip("#"))
            title = strip_md_links(line[hashes:].strip())
            if hashes <= 1:
                heading1(doc, title)
            elif hashes == 2:
                heading1(doc, title)
            else:
                heading2(doc, title)
            i += 1
            continue
        if line.lstrip().startswith(">"):
            quote = strip_md_links(line.lstrip()[1:].strip())
            p = doc.add_paragraph()
            p.paragraph_format.space_after = Pt(8)
            p.paragraph_format.left_indent = Cm(0.5)
            add_md_runs(p, quote, size=11, color=MUTED)
            for run in p.runs:
                run.italic = True
            i += 1
            continue
        m_num = re.match(r"^(\d+)\.\s+(.*)$", line.strip())
        if m_num:
            p = doc.add_paragraph(style="List Number")
            p.paragraph_format.space_after = Pt(4)
            add_md_runs(p, m_num.group(2))
            i += 1
            continue
        if re.match(r"^[-*]\s+", line.strip()):
            text = re.sub(r"^[-*]\s+", "", line.strip())
            p = doc.add_paragraph(style="List Bullet")
            p.paragraph_format.space_after = Pt(4)
            add_md_runs(p, text)
            i += 1
            continue
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(8)
        p.paragraph_format.line_spacing = 1.15
        add_md_runs(p, line.strip())
        i += 1
    if in_code:
        flush_code()


def header_footer(doc):
    section = doc.sections[0]
    header = section.header
    hp = header.paragraphs[0]
    hp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    r = hp.add_run("Smart Agro Community  ·  Competition Proposal")
    set_run_font(r, size=9, italic=True, color=MUTED)

    footer = section.footer
    fp = footer.paragraphs[0]
    fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = fp.add_run("Confidential — for judging panel  ·  © 2026  ·  Page ")
    set_run_font(r, size=9, color=MUTED)
    # PAGE field
    fld = OxmlElement("w:fldChar")
    fld.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = " PAGE "
    fld2 = OxmlElement("w:fldChar")
    fld2.set(qn("w:fldCharType"), "end")
    run = fp.add_run()
    run._r.append(fld)
    run2 = fp.add_run()
    run2._r.append(instr)
    run3 = fp.add_run()
    run3._r.append(fld2)
    set_run_font(run, size=9, color=MUTED)
    set_run_font(run2, size=9, color=MUTED)
    set_run_font(run3, size=9, color=MUTED)


def build():
    doc = Document()
    section = doc.sections[0]
    section.page_width = Inches(8.27)
    section.page_height = Inches(11.69)
    section.left_margin = Cm(2.2)
    section.right_margin = Cm(2.2)
    section.top_margin = Cm(2.0)
    section.bottom_margin = Cm(2.0)
    header_footer(doc)

    # Cover block
    para(doc, "WEB PROJECT COMPETITION PROPOSAL", size=11, bold=True, color=GREEN2, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=4)
    p = para(doc, "Smart Agro Community", size=28, bold=True, color=GREEN, align=WD_ALIGN_PARAGRAPH.CENTER, space_before=6, space_after=4)
    add_bottom_border(p, sz="18")
    para(
        doc,
        "Smart Tools for Myanmar Farmers",
        size=16,
        italic=True,
        color=GREEN2,
        align=WD_ALIGN_PARAGRAPH.CENTER,
        space_before=8,
        space_after=14,
    )
    para(
        doc,
        "A bilingual agricultural intelligence platform that turns a phone photo into a field diagnosis, a trusted advisor into a conversation, and isolated farms into a connected community.",
        size=12,
        italic=True,
        color=INK,
        align=WD_ALIGN_PARAGRAPH.CENTER,
        space_after=16,
    )

    meta = [
        ("Project type", "Web application (Progressive Web App)"),
        ("Domain", "AgriTech  ·  Digital inclusion  ·  Climate-smart farming"),
        ("Primary users", "Myanmar farmers, agronomy experts, and platform administrators"),
        ("Languages", "English and Myanmar (မြန်မာ)"),
        ("Tagline", "Smart Tools for Myanmar Farmers"),
        ("Year", "2026"),
    ]
    add_table(doc, ["Item", "Detail"], meta)

    heading1(doc, "Team Members")
    para(
        doc,
        "The project is submitted by the following team from the University of Computer Studies (Meiktila).",
        space_after=8,
    )
    members = add_table(
        doc,
        ["No.", "Name", "Student ID"],
        [
            ["1", "Mg Arkar Thet Naing", "24-25-UCSMTLA-"],
            ["2", "Mg Khant Zaw", "24-25-UCSMTLA-"],
            ["3", "Mg Kaung Myat Tun", "24-25-UCSMTLA-"],
            ["4", "Mg Yawai Aung", "24-25-UCSMTLA-"],
        ],
    )
    set_cell_widths(members, [1.8, 8.2, 6.0])

    heading1(doc, "Official Project Evaluation Criteria")
    para(
        doc,
        "The photograph below is the University of Computer Studies (Meiktila) 2025–2026 / 2026–2027 Academic Year Project Evaluation Criteria. It is attached here as an official source document for this submission.",
    )
    if EVAL_IMG.exists():
        add_picture(doc, EVAL_IMG, width_in=6.1)
        para(
            doc,
            "Figure 1. Official evaluation criteria (University of Computer Studies, Meiktila).",
            size=10,
            italic=True,
            color=MUTED,
            align=WD_ALIGN_PARAGRAPH.CENTER,
            space_after=12,
        )
    else:
        para(doc, "[Evaluation-criteria photograph missing from docs folder.]", italic=True, color=MUTED)

    para(
        doc,
        "Typed transcription of the attached criteria (for readability if the photograph is reduced in print):",
        space_before=4,
    )
    add_table(
        doc,
        ["No.", "Criterion", "Marks", "What is assessed"],
        [
            [
                "1",
                "Requirements and System Analysis",
                "15",
                "Problem definition (agriculture / livestock / SME pain points); feasibility for smallholders; system design (UML, ER).",
            ],
            [
                "2",
                "Technical Implementation & Features",
                "20",
                "UI/UX and accessibility for rural users (mobile-first / simplified inputs); adherence to the SRS functional and non-functional requirements.",
            ],
            [
                "3",
                "Data Analytics & Reporting",
                "20",
                "Accurate charts, tables, and KPIs; decision support (actionable insights and automated recommendations).",
            ],
            [
                "4",
                "Software Quality & Architecture",
                "15",
                "Modular code, documentation, testing, security/privacy, and live-deployment readiness.",
            ],
            [
                "5",
                "Domain Relevance & Business Viability",
                "30",
                "Innovative support for agricultural SMEs; cost–benefit; local industry applicability; quality of report and defence.",
            ],
            ["", "Total", "100", ""],
        ],
    )
    heading2(doc, "How this submission maps to the criteria")
    add_table(
        doc,
        ["Criterion", "Where it is evidenced in this document"],
        [
            [
                "1. Requirements & analysis (15)",
                "Problem Statement (§2); feasibility in Technology Stack (§4); full Software Requirements Specification attached as Appendix A (scope, users, constraints, data models, APIs).",
            ],
            [
                "2. Technical implementation (20)",
                "Core Solutions (§3): Detect, BaGyi Pyoe, Community, weather, heatmap, Knowledge Center; bilingual Myanmar-first PWA for phone users.",
            ],
            [
                "3. Data analytics & reporting (20)",
                "Detection history and bilingual lab reports; outbreak heatmap; weather-linked advice; admin/expert queues as operational KPIs.",
            ],
            [
                "4. Quality & architecture (15)",
                "Typed Express API, Zod validation, Helmet/CORS/rate limits, JWT+OTP, Vitest tests; architecture in §4 and SRS §6–10.",
            ],
            [
                "5. Domain & viability (30)",
                "Social Impact (§5): same-day diagnosis, IPM-first advice, reduced crop loss; Myanmar crops and DOA protocols. Full SRS attached for documentation defence.",
            ],
        ],
    )
    para(
        doc,
        "The complete Software Requirements Specification (SRS v4.1) is attached as Appendix A at the end of this proposal — not only cited by filename.",
        italic=True,
        color=MUTED,
    )

    # 1
    heading1(doc, "1. Executive Summary & Project Vision")
    para(
        doc,
        "Agriculture remains the backbone of Myanmar’s rural economy. Yet the farmer who notices a yellowing leaf at dawn still faces a familiar delay: wait for an extension officer, guess at a pesticide, or lose a season to a disease that could have been named in minutes.",
    )
    mixed_para(
        doc,
        [
            ("Smart Agro Community", True, False),
            (" closes that gap. It is a full-stack, mobile-first web platform that puts three capabilities in one place:", False, False),
        ],
    )
    bullet(doc, " — AI disease and pest detection from a field photo, across 19 Myanmar crops.", bold_lead="See")
    bullet(doc, " — BaGyi Pyoe / ဘကြီးပျိုး, a weather-aware farming chatbot that answers in the farmer’s own language.", bold_lead="Ask")
    bullet(doc, " — an interactive community where diagnoses, weather, and local experience travel faster than an outbreak.", bold_lead="Share")
    mixed_para(
        doc,
        [
            ("The vision is not “another dashboard.” It is a ", False, False),
            ("field companion", True, False),
            (": Myanmar-first in language and treatment advice, grounded in Department of Agriculture (DOA) field guides, and designed so a farmer on a phone can act the same day.", False, False),
        ],
    )
    para(
        doc,
        "We built a complete product, not a prototype slide. Detection, chat, community, weather, outbreak mapping, expert review, knowledge resources, and admin moderation are live in the codebase and connected through a single farmer experience.",
    )
    q = para(
        doc,
        "If a farmer can photograph a leaf, they should be able to name the problem, know what to do, and learn from neighbours — in Myanmar, on a phone, today.",
        size=12,
        italic=True,
        bold=True,
        color=GREEN,
        align=WD_ALIGN_PARAGRAPH.CENTER,
        space_before=10,
        space_after=14,
    )

    # 2
    heading1(doc, "2. Problem Statement")
    heading2(doc, "2.1 Delayed diagnosis costs harvests")
    para(
        doc,
        "Crop disease and pest pressure move faster than traditional advisory channels. By the time a farmer reaches a township office or waits for a visit, infection may already be in neighbouring plots. Misidentification leads to the wrong chemical, wasted money, and preventable yield loss.",
    )
    heading2(doc, "2.2 Advice is fragmented and language-barred")
    para(
        doc,
        "Field knowledge exists — in DOA manuals, in expert heads, in neighbouring farms — but it is not where the farmer is. English-only apps, PDF manuals, and generic global models fail Myanmar’s script, crops, and cropping calendar. A rice blast protocol copied from another country is not a Myanmar paddy protocol.",
    )
    heading2(doc, "2.3 Isolation during outbreaks")
    para(
        doc,
        "Farmers often discover the same pest in the same week without knowing it is a regional pattern. There is no shared, township-level picture of what is being detected today, and no trusted way to ask “has anyone else seen this?” with a photo and a verified diagnosis attached.",
    )
    heading2(doc, "2.4 Trust and safety on digital platforms")
    para(
        doc,
        "A community without moderation becomes noise or harm. Reports need a reason, owners need to delete their own posts, and admins need to approve or deny removals with a notice back to the author. Expert review of AI results is equally essential: models assist; agronomists confirm.",
    )
    mixed_para(
        doc,
        [
            ("Smart Agro Community is built for these four problems together", True, False),
            (" — diagnosis, advice, community signal, and accountable review — rather than as disconnected tools.", False, False),
        ],
    )

    # 3
    heading1(doc, "3. Core Solutions")
    para(
        doc,
        "The platform is organised around three pillars requested by this competition, implemented as they exist in the product today.",
    )

    heading2(doc, "3.1 AI Detection — from photo to field protocol")
    mixed_para(
        doc,
        [
            ("What the farmer does: ", True, False),
            ("Opens Detect, optionally sets township or GPS, uploads a clear photo of a leaf, stem, fruit, or pest damage (JPEG / PNG / WebP, up to 5 MB), and receives a structured result.", False, False),
        ],
    )
    para(doc, "What the system does:", bold=True, space_after=6)
    add_table(
        doc,
        ["Capability", "Implementation in product"],
        [
            [
                "Multi-crop recognition",
                "19 field crops: Rice, Black Gram, Green Gram, Pigeon Pea, Sesame, Groundnut, Sunflower, Maize, Rubber, Sugarcane, Cotton, Cabbage, Onion, Garlic, Chili, Potato, Tea, Coffee, Oil Palm",
            ],
            [
                "Disease & pest catalog",
                "Canonical English labels plus Myanmar names; hundreds of unique problems including Healthy",
            ],
            [
                "Vision pipeline",
                "Google Gemini vision as primary (API-key rotation on quota); Cursor as fallback; optional local scikit-learn SVM rice model (HSV + HOG + LBP + GLCM)",
            ],
            ["Quality gate", "Rejects non-crop / non-leaf-like images rather than inventing a disease"],
            [
                "Result card",
                "Crop, disease (Myanmar-first), confidence, severity (Mild → Critical), ranked alternatives",
            ],
            [
                "Treatment",
                "DOA Myanmar field protocols (cotton, pulses, oilseeds, pepper/chili, rice IPM) preferred over generic AI text",
            ],
            [
                "Expert loop",
                "Farmer can request expert review; experts/admins verify, correct, or reject with a reason; farmer receives a notice",
            ],
            [
                "Lab report",
                "Downloadable bilingual .docx laboratory-style report (Myanmar before English)",
            ],
            [
                "Outbreak signal",
                "Each detection can carry location and weather context and feed the community heatmap",
            ],
        ],
    )
    para(
        doc,
        "Detection is not a black box. The farmer sees confidence, can expand a full treatment guide (symptoms, cultural control, recommended chemistry with label-use caution), share to Community, or ask an expert. That is how AI earns trust in the field.",
    )

    heading2(doc, "3.2 AI Chatbot — BaGyi Pyoe / ဘကြီးပျိုး")
    para(doc, "The assistant is not a generic “bot.” It is a named field elder:")
    bullet(doc, " BaGyi Pyoe", bold_lead="English:")
    bullet(doc, " ဘကြီးပျိုး", bold_lead="Myanmar:")
    para(
        doc,
        "The name is intentional: a trusted uncle who raises seedlings — warm, practical, never a cold chatbot brand.",
    )
    para(doc, "What it does in the product:", bold=True, space_after=4)
    bullet(doc, "Answers in the same language the farmer uses (English or Myanmar).")
    bullet(doc, "Knows the supported crop list and crop-specific disease/pest problems.")
    bullet(doc, "Receives live weather context (temperature, humidity, rain, alerts) for the farmer’s township or GPS point and must use those numbers — not invent Yangon weather.")
    bullet(doc, "Ties weather to risk (e.g. humidity + rain → fungal pressure; stagnant water + high nitrogen → planthopper).")
    bullet(doc, "Prefers IPM: scouting, sanitation, resistant varieties; does not invent unverified brand dosages.")
    bullet(doc, "Saves chat history for signed-in users; guests can try the advisor without an account.")
    bullet(doc, "Offers quick suggestions (blast care, pest IPM, local forecast).")
    para(
        doc,
        "This is agricultural decision support, not entertainment chat: short steps, what to check in the field today, and when to open Detect or Weather instead of guessing.",
    )

    heading2(doc, "3.3 Interactive Community — learn together, stay safe")
    para(doc, "Community is a social layer built for farms, not for generic social media.")
    para(doc, "Feed", bold=True, color=GREEN2, space_after=4)
    bullet(doc, "Field posts with text and photos; optional link to a verified diagnosis.")
    bullet(doc, "Likes, comments, and nested replies.")
    bullet(doc, "Save posts locally; add friend / message the author from the card.")
    bullet(doc, "Author can delete their own post.")
    bullet(doc, "Others can report a post with a structured reason (spam, harassment, false information, inappropriate, other + details).")
    para(doc, "Trust & moderation", bold=True, color=GREEN2, space_after=4)
    bullet(doc, "Reports go to admins (inbox + Community moderation queue).")
    bullet(doc, "Admin approves (deletes the post; author receives a notice with the deletion reason) or denies (post stays).")
    bullet(doc, "Hide / restore remains available for broader moderation.")
    bullet(doc, "Farmers can appeal certain notices.")
    para(doc, "Beyond the feed", bold=True, color=GREEN2, space_after=4)
    bullet(doc, " one-to-one and group chats, friend requests, invite links.", bold_lead="Messages:")
    bullet(doc, " farm crops, township, bio, own posts.", bold_lead="Profiles:")
    bullet(doc, " books, articles, and journals for practical reading.", bold_lead="Knowledge Center:")
    bullet(doc, " Open-Meteo current conditions and 7-day forecast, township search, GPS, crop-related tips.", bold_lead="Weather:")
    bullet(doc, " Leaflet choropleth + heat layer of community detections across Myanmar townships and regions.", bold_lead="Outbreak heatmap:")
    bullet(doc, " Farmer · Expert · Admin — experts review diagnoses; admins run users, knowledge, moderation, and audit.", bold_lead="Roles:")
    para(
        doc,
        "Together, Community turns a private scan into public early warning — with the safeguards a public agricultural network requires.",
    )

    # 4
    heading1(doc, "4. Technology Stack")
    para(doc, "The stack below is what the repository actually runs — not a hypothetical architecture.")

    heading2(doc, "4.1 Client (farmer & staff web app)")
    add_table(
        doc,
        ["Layer", "Technology"],
        [
            ["UI", "React 19, TypeScript, React Router 7"],
            ["Build", "Vite 6, @vitejs/plugin-react"],
            ["Installable app", "vite-plugin-pwa (Workbox, standalone manifest: Smart Agro Community)"],
            ["Maps", "Leaflet, leaflet.heat, Myanmar GeoJSON boundaries"],
            ["Documents", "docx (client-generated lab reports)"],
            ["i18n", "First-class English / Myanmar copy, Myanmar numerals, Padauk & Noto Sans Myanmar"],
            ["Styling", "Custom design system (mobile-first, Myanmar line-height and wrap rules)"],
        ],
    )

    heading2(doc, "4.2 API server")
    add_table(
        doc,
        ["Layer", "Technology"],
        [
            ["Runtime", "Node.js, Express 4, TypeScript (tsx in development)"],
            ["Database", "MongoDB via Mongoose 8"],
            ["Files", "GridFS (default); S3-ready configuration"],
            ["Auth", "JWT access + refresh, bcryptjs, OTP rate limits"],
            ["Validation", "Zod on request bodies"],
            ["Security", "Helmet, CORS, express-rate-limit, magic-number checks on uploads (file-type, multer)"],
            ["Mail", "Nodemailer"],
            ["Tests", "Vitest, Supertest, mongodb-memory-server"],
        ],
    )
    mixed_para(
        doc,
        [
            ("API surface (representative): ", True, False),
            ("/auth, /detections, /chatbot, /social, /messages, /weather, /heatmap, /knowledge, /admin, /files.", False, False),
        ],
    )

    heading2(doc, "4.3 Intelligence layer")
    add_table(
        doc,
        ["Layer", "Technology"],
        [
            ["Primary vision & chat", "Google Gemini (configurable model, e.g. Gemini 3.6 Flash; multi-key rotation)"],
            ["Fallback LLM / agent", "Cursor SDK (@cursor/sdk)"],
            ["Local rice model (optional)", "Python FastAPI — OpenCV, Pillow, scikit-learn SVM, scikit-image (HOG / LBP / GLCM)"],
            ["Weather", "Open-Meteo forecast + geocoding (Myanmar-restricted search)"],
            ["Field knowledge", "Curated DOA Myanmar treatment catalog mapped to detect labels"],
        ],
    )

    heading2(doc, "4.4 Architecture in one sentence")
    mixed_para(
        doc,
        [
            ("A ", False, False),
            ("React PWA", True, False),
            (" talks to a ", False, False),
            ("typed Express API", True, False),
            (" on MongoDB; ", False, False),
            ("Gemini", True, False),
            (" (with Cursor fallback) diagnoses photos and powers BaGyi Pyoe; ", False, False),
            ("Open-Meteo", True, False),
            (" grounds advice in real township weather; ", False, False),
            ("DOA protocols", True, False),
            (" and expert review keep recommendations accountable.", False, False),
        ],
    )
    para(doc, "System flow", bold=True, space_after=4, color=GREEN2)
    for line in [
        "[ Phone / Browser PWA ]",
        "              │",
        "              ▼",
        "[ Vite + React + TypeScript ]",
        "              │  /api/v1",
        "              ▼",
        "[ Express + JWT + Zod + Helmet ]",
        "              │",
        "              ├── MongoDB + GridFS",
        "              ├── Gemini vision & chat  →  Cursor fallback",
        "              ├── FastAPI rice SVM (optional)",
        "              └── Open-Meteo + DOA field guides",
    ]:
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(0)
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.line_spacing = 1.0
        run = p.add_run(line)
        set_run_font(run, name="Consolas", size=10, color=GREEN)

    # 5
    heading1(doc, "5. Social Impact & Future Roadmap")
    heading2(doc, "5.1 Who benefits")
    add_table(
        doc,
        ["Stakeholder", "Impact"],
        [
            ["Smallholder farmers", "Same-day, Myanmar-language diagnosis and IPM steps without waiting for a distant office"],
            ["Women and youth on farms", "Phone-first access; no requirement for English literacy"],
            ["Extension / experts", "A queue of real photos and cases instead of unstructured chat rumours"],
            ["Townships & DOA-aligned practice", "Advice drawn from national field guides, not anonymous internet recipes"],
            ["Communities", "Heatmap and feed make outbreaks visible while they can still be contained"],
        ],
    )

    heading2(doc, "5.2 Why this is more than an app")
    para(
        doc,
        "Food security in Myanmar is a field problem: humidity, monsoon timing, seed quality, and neighbour-to-neighbour spread. Digital tools fail when they ignore language, local crops, and trust. Smart Agro Community is localised by design — 19 crops that are actually grown here, Myanmar-first UI and reports, township weather, and a character (ဘကြီးပျိုး) farmers can talk to as an elder, not as a vendor.",
    )
    para(
        doc,
        "Responsible AI is built in: quality gates, confidence scores, expert override, pesticide-label caution, and community reporting with human approval before a post is deleted.",
    )

    heading2(doc, "5.3 Alignment with competition values")
    bullet(doc, " Multi-crop vision + weather-aware chat + outbreak map in one bilingual PWA.", bold_lead="Innovation:")
    bullet(doc, " Running MERN-style stack, real APIs, PWA install, and documented local services.", bold_lead="Feasibility:")
    bullet(doc, " Myanmar script, digits, and DOA protocols; guest try-before-register for chat.", bold_lead="Inclusion:")
    bullet(doc, " IPM-first advice reduces panic spraying; community signal reduces duplicate loss.", bold_lead="Sustainability:")
    bullet(doc, " Crop catalog and treatment maps are data, not hardcoded one-crop logic.", bold_lead="Scalability:")

    heading2(doc, "5.4 Future roadmap")
    para(doc, "Near term", bold=True, space_after=4)
    bullet(doc, "More DOA crop manuals in the treatment catalog (vegetables, plantation crops).")
    bullet(doc, "Offline-friendly caching of saved reports and recently viewed guides.")
    bullet(doc, "Stronger push notifications for outbreak clusters in the farmer’s township.")
    para(doc, "Medium term", bold=True, space_before=8, space_after=4)
    bullet(doc, "On-device or edge models for rice and chili where connectivity is weak.")
    bullet(doc, "Partnership pilots with township agriculture offices using the expert-review queue.")
    bullet(doc, "Voice input for BaGyi Pyoe for low-literacy users.")
    para(doc, "Long term", bold=True, space_before=8, space_after=4)
    bullet(doc, "Seasonal risk forecasts combining heatmap history and weather.")
    bullet(doc, "Open data dashboards for researchers (privacy-preserving, township-level).")
    bullet(doc, "A national early-warning layer: when many farmers scan the same pest, the map speaks before the crop fails.")

    heading1(doc, "Closing statement")
    para(doc, "Smart Agro Community is a working answer to a simple, urgent question:")
    para(
        doc,
        "When a leaf changes colour in a Myanmar field, how fast can the farmer know, act, and warn the next farm?",
        size=13,
        bold=True,
        italic=True,
        color=GREEN,
        align=WD_ALIGN_PARAGRAPH.CENTER,
        space_before=8,
        space_after=10,
    )
    mixed_para(
        doc,
        [
            ("We built the camera, the advisor, and the community so that answer can be: ", False, False),
            ("this morning.", True, True, GREEN),
        ],
        align=WD_ALIGN_PARAGRAPH.CENTER,
    )
    para(
        doc,
        "We respectfully submit Smart Agro Community for your consideration — a complete, bilingual, field-ready platform worthy of first place.",
        align=WD_ALIGN_PARAGRAPH.CENTER,
        space_before=8,
    )
    p = para(doc, "Smart Agro Community", size=14, bold=True, color=GREEN, align=WD_ALIGN_PARAGRAPH.CENTER, space_before=18, space_after=2)
    para(doc, "Smart Tools for Myanmar Farmers", size=11, italic=True, color=GREEN2, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=2)
    para(doc, "© 2026  ·  Built for Myanmar farms", size=10, color=MUTED, align=WD_ALIGN_PARAGRAPH.CENTER)

    page_break(doc)
    heading1(doc, "Appendix A — Software Requirements Specification (full text)")
    para(
        doc,
        "This appendix reproduces the authoritative SRS for Smart Agro Community (document version 4.1). It is attached in full so evaluators do not need a separate file to read the requirements, interfaces, data models, testing, and deployment notes.",
        space_after=12,
    )
    if SRS_MD.exists():
        append_markdown(doc, SRS_MD.read_text(encoding="utf-8"))
    else:
        para(doc, "[SRS markdown file was not found.]", italic=True, color=MUTED)

    saved = save_docx(doc, OUT)
    print("proposal:", saved)

    if SRS_MD.exists():
        srs_doc = Document()
        s = srs_doc.sections[0]
        s.page_width = Inches(8.27)
        s.page_height = Inches(11.69)
        s.left_margin = Cm(2.2)
        s.right_margin = Cm(2.2)
        s.top_margin = Cm(2.0)
        s.bottom_margin = Cm(2.0)
        header_footer(srs_doc)
        append_markdown(srs_doc, SRS_MD.read_text(encoding="utf-8"))
        srs_saved = save_docx(srs_doc, SRS_OUT)
        print("srs:", srs_saved)


if __name__ == "__main__":
    build()
