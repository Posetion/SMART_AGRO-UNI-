"""Build a fill-in Word template for the Smart Agro competition entry."""
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Inches, Pt, RGBColor

OUT = Path(r"d:\SMART-AGRO\docs\Smart-Agro-Community-Competition-Entry.docx")
GREEN = RGBColor(0x1B, 0x43, 0x32)
GREEN2 = RGBColor(0x2E, 0x7D, 0x32)
INK = RGBColor(0x21, 0x21, 0x21)
MUTED = RGBColor(0x5A, 0x65, 0x5C)
WHITE = RGBColor(0xFF, 0xFF, 0xFF)
HEADER_BG = "1B4332"
ROW_ALT = "F3F8F4"
LINE = "B7C9BC"


def set_cell_shading(cell, color: str):
    tc_pr = cell._tc.get_or_add_tcPr()
    shading = OxmlElement("w:shd")
    shading.set(qn("w:fill"), color)
    tc_pr.append(shading)


def set_cell_border(cell, color=LINE, size="8"):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    borders = tc_pr.first_child_found_in("w:tcBorders")
    if borders is None:
        borders = OxmlElement("w:tcBorders")
        tc_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        tag = "w:" + edge
        element = borders.find(qn(tag))
        if element is None:
            element = OxmlElement(tag)
            borders.append(element)
        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), size)
        element.set(qn("w:color"), color)


def set_font(run, size=10.5, bold=False, italic=False, color=INK):
    run.font.name = "Calibri"
    run._element.get_or_add_rPr().rFonts.set(qn("w:eastAsia"), "Myanmar Text")
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic
    run.font.color.rgb = color


def paragraph(container, text="", size=10.5, bold=False, italic=False, color=INK, align=None, after=4):
    p = container.add_paragraph()
    p.paragraph_format.space_after = Pt(after)
    p.paragraph_format.line_spacing = 1.08
    if align is not None:
        p.alignment = align
    if text:
        set_font(p.add_run(text), size=size, bold=bold, italic=italic, color=color)
    return p


def label_value_table(doc):
    rows = [
        ("Team leader name", "အဖွဲ့ခေါင်းဆောင်အမည်", "[Fill in team leader name]"),
        ("Team leader email", "အဖွဲ့ခေါင်းဆောင် Email", "[Fill in email address]"),
        ("Project title", "စီမံကိန်းအမည်", "Smart Agro Community"),
    ]
    table = doc.add_table(rows=0, cols=3)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    widths = [Cm(4.4), Cm(4.4), Cm(9.2)]
    for index, (en, my, value) in enumerate(rows):
        cells = table.add_row().cells
        for cell, width in zip(cells, widths):
            cell.width = width
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            set_cell_border(cell)
        set_cell_shading(cells[0], ROW_ALT)
        set_cell_shading(cells[1], ROW_ALT)
        paragraph(cells[0], en, size=10, bold=True, color=GREEN)
        paragraph(cells[1], my, size=10, bold=True, color=GREEN)
        paragraph(cells[2], value, size=10.5, bold=value == "Smart Agro Community", color=INK)
    return table


def text_section(doc, title, my_title, limit, text):
    paragraph(doc, f"{my_title} / {title}  (Maximum {limit} words)", size=12, bold=True, color=GREEN2, after=3)
    box = doc.add_table(rows=1, cols=1)
    box.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell = box.cell(0, 0)
    cell.width = Cm(18)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.TOP
    set_cell_border(cell, color=LINE, size="10")
    paragraph(cell, text, size=10.5, color=INK, after=4)
    return box


def image_placeholder(doc, title, subtitle, width_cm=8.7, height_cm=5.0):
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell = table.cell(0, 0)
    cell.width = Cm(width_cm)
    cell.height = Cm(height_cm)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    set_cell_border(cell, color=GREEN2, size="12")
    set_cell_shading(cell, "F7FBF8")
    paragraph(cell, title, size=11, bold=True, color=GREEN, align=WD_ALIGN_PARAGRAPH.CENTER, after=5)
    paragraph(cell, subtitle, size=9, italic=True, color=MUTED, align=WD_ALIGN_PARAGRAPH.CENTER, after=0)
    return table


def add_page_number(section):
    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = footer.add_run("Smart Agro Community  |  Competition Entry")
    set_font(run, size=8.5, color=MUTED)


def main():
    doc = Document()
    section = doc.sections[0]
    section.top_margin = Cm(1.4)
    section.bottom_margin = Cm(1.4)
    section.left_margin = Cm(1.6)
    section.right_margin = Cm(1.6)
    add_page_number(section)

    p = paragraph(doc, "PROJECT COMPETITION ENTRY", size=10, bold=True, color=GREEN2, align=WD_ALIGN_PARAGRAPH.CENTER, after=3)
    p = paragraph(doc, "Smart Agro Community", size=24, bold=True, color=GREEN, align=WD_ALIGN_PARAGRAPH.CENTER, after=3)
    paragraph(doc, "Smart Tools for Myanmar Farmers", size=12, italic=True, color=MUTED, align=WD_ALIGN_PARAGRAPH.CENTER, after=12)

    paragraph(doc, "Project Information / စီမံကိန်းအချက်အလက်", size=14, bold=True, color=GREEN, after=6)
    label_value_table(doc)
    paragraph(doc, "", after=2)

    text_section(
        doc,
        "Description",
        "ဖော်ပြချက်",
        30,
        "A bilingual mobile-first platform that helps Myanmar farmers identify crop diseases, receive practical advice, share local knowledge, and respond quickly to weather and outbreak risks.",
    )
    paragraph(doc, "", after=2)
    text_section(
        doc,
        "Objective",
        "ရည်ရွယ်ချက်",
        20,
        "Help Myanmar farmers identify crop diseases, receive practical advice, share local knowledge, and respond quickly to weather and outbreak risks.",
    )
    paragraph(doc, "", after=2)
    text_section(
        doc,
        "Abstract",
        "အကျဉ်းချုပ်",
        30,
        "Smart Agro Community combines AI crop diagnosis, a weather-aware farming assistant, community discussion, expert review, and outbreak mapping in one accessible bilingual web application.",
    )

    paragraph(doc, "", after=2)
    paragraph(doc, "Images / ပုံများ", size=14, bold=True, color=GREEN, after=5)
    image_table = doc.add_table(rows=1, cols=2)
    image_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    image_table.autofit = False
    for cell, title, subtitle in zip(
        image_table.rows[0].cells,
        ("Cover image / ပင်မပုံ", "Project image 1 / စီမံကိန်းပုံ ၁"),
        ("Insert 1 image", "Maximum 1 MB"),
    ):
        cell.width = Cm(8.7)
        set_cell_border(cell, color="2E7D32", size="12")
        set_cell_shading(cell, "F7FBF8")
        paragraph(cell, title, size=10.5, bold=True, color=GREEN, align=WD_ALIGN_PARAGRAPH.CENTER, after=4)
        paragraph(cell, subtitle, size=9, italic=True, color=MUTED, align=WD_ALIGN_PARAGRAPH.CENTER, after=0)
    for image_number in (2, 3):
        row = image_table.add_row()
        for cell in row.cells:
            cell.width = Cm(8.7)
            set_cell_border(cell, color="2E7D32", size="12")
            set_cell_shading(cell, "F7FBF8")
        paragraph(row.cells[0], f"Project image {image_number} / စီမံကိန်းပုံ {image_number}", size=10.5, bold=True, color=GREEN, align=WD_ALIGN_PARAGRAPH.CENTER, after=4)
        paragraph(row.cells[0], "Insert 1 image · Maximum 1 MB", size=9, italic=True, color=MUTED, align=WD_ALIGN_PARAGRAPH.CENTER, after=0)
        paragraph(row.cells[1], f"Project image {image_number + 1} / စီမံကိန်းပုံ {image_number + 1}", size=10.5, bold=True, color=GREEN, align=WD_ALIGN_PARAGRAPH.CENTER, after=4)
        paragraph(row.cells[1], "Insert 1 image · Maximum 1 MB", size=9, italic=True, color=MUTED, align=WD_ALIGN_PARAGRAPH.CENTER, after=0)

    doc.add_page_break()
    paragraph(doc, "Team Members / အဖွဲ့ဝင်များ", size=16, bold=True, color=GREEN, after=4)
    paragraph(doc, "Complete the table with each member's name, age in years, and major.", size=10, italic=True, color=MUTED, after=8)
    table = doc.add_table(rows=1, cols=4)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    headers = ["No.", "Name / အမည်", "Years / အသက်", "Major / အထူးပြုဘာသာ"]
    widths = [Cm(1.5), Cm(6.2), Cm(3.2), Cm(7.1)]
    for cell, header, width in zip(table.rows[0].cells, headers, widths):
        cell.width = width
        set_cell_border(cell, color=HEADER_BG, size="10")
        set_cell_shading(cell, HEADER_BG)
        paragraph(cell, header, size=10, bold=True, color=WHITE, align=WD_ALIGN_PARAGRAPH.CENTER, after=0)
    for number in range(1, 7):
        cells = table.add_row().cells
        for cell, width in zip(cells, widths):
            cell.width = width
            cell.height = Cm(0.85)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            set_cell_border(cell)
        if number % 2 == 0:
            for cell in cells:
                set_cell_shading(cell, ROW_ALT)
        paragraph(cells[0], str(number), size=10.5, bold=True, color=GREEN, align=WD_ALIGN_PARAGRAPH.CENTER, after=0)
        paragraph(cells[1], "", after=0)
        paragraph(cells[2], "", after=0)
        paragraph(cells[3], "", after=0)

    paragraph(doc, "", after=4)
    paragraph(doc, "Submission checklist / တင်သွင်းရန် စစ်ဆေးရန်", size=14, bold=True, color=GREEN, after=5)
    checklist = [
        "Team leader name and email completed",
        "Description: maximum 30 words",
        "Objective: maximum 20 words",
        "Abstract: maximum 30 words",
        "One cover image attached, maximum 1 MB",
        "Three project images attached, maximum 1 MB each",
        "Team member table completed",
    ]
    for item in checklist:
        p = doc.add_paragraph(style="List Bullet")
        p.paragraph_format.space_after = Pt(3)
        set_font(p.add_run(item), size=10, color=INK)

    doc.core_properties.title = "Smart Agro Community Competition Entry"
    doc.core_properties.subject = "Project competition submission template"
    doc.core_properties.author = "Smart Agro Community"
    OUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUT)
    print(OUT)


if __name__ == "__main__":
    main()
